import streamlit as st
import sqlite3
from datetime import datetime
from fpdf import FPDF
from docx import Document

# Conexão com o banco
def conectar():
    return sqlite3.connect('ocorrencias.db')

# Criação das tabelas
def criar_tabelas():
    conn = conectar()
    cursor = conn.cursor()

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS usuarios (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nome TEXT NOT NULL,
        usuario TEXT UNIQUE NOT NULL,
        senha TEXT NOT NULL,
        setor TEXT NOT NULL
    )
    """)

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS alunos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        cgm TEXT UNIQUE NOT NULL,
        nome TEXT NOT NULL,
        data_nascimento TEXT,
        telefone TEXT,
        responsavel TEXT,
        data TEXT,
        turma TEXT
    )
    """)

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS ocorrencias (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        cgm TEXT NOT NULL,
        nome TEXT,
        telefone TEXT,
        data TEXT,
        descricao TEXT
    )
    """)
    conn.commit()
    conn.close()

# Login
def login():
    st.title("Login 👤")
    usuario = st.text_input("Usuário")
    senha = st.text_input("Senha", type="password")
    if st.button("Entrar"):
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM usuarios WHERE usuario=? AND senha=?", (usuario, senha))
        resultado = cursor.fetchone()
        conn.close()
        if resultado:
            st.session_state["logado"] = True
            st.rerun()
        else:
            st.error("Usuário ou senha inválidos!")

# Página Cadastro de Alunos
def pagina_cadastro_aluno():
    st.title("Cadastro de Alunos 👦👧")

    arquivo = st.file_uploader("Importar lista de alunos (.txt)", type="txt")
    if arquivo:
        conteudo = arquivo.read().decode('utf-8')
        linhas = conteudo.strip().split('\n')[1:]
        dados_alunos = []
        for linha in linhas:
            campos = linha.split('\t')
            if len(campos) >= 3:
                cgm, nome, telefone = campos[0], campos[1], campos[2]
                dados_alunos.append((cgm, nome, telefone))

        if dados_alunos:
            st.success(f"{len(dados_alunos)} alunos carregados.")
            st.dataframe(dados_alunos)

            if st.button("Salvar no banco"):
                conn = conectar()
                cursor = conn.cursor()
                for aluno in dados_alunos:
                    try:
                        cursor.execute("""
                        INSERT INTO alunos (cgm, nome, telefone)
                        VALUES (?, ?, ?)
                        """, aluno)
                    except sqlite3.IntegrityError:
                        pass
                conn.commit()
                conn.close()
                st.success("Alunos salvos com sucesso!")

    st.subheader("Cadastrar aluno manualmente")
    cgm = st.text_input("CGM")
    nome = st.text_input("Nome")
    data_nascimento = st.date_input("Data de Nascimento")
    telefone = st.text_input("Telefone")
    responsavel = st.text_input("Responsável")
    data = st.date_input("Data de Cadastro")
    turma = st.text_input("Turma")

    if st.button("Cadastrar Aluno"):
        conn = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                INSERT INTO alunos (cgm, nome, data_nascimento, telefone, responsavel, data, turma)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (cgm, nome, str(data_nascimento), telefone, responsavel, str(data), turma))
            conn.commit()
            st.success("Aluno cadastrado com sucesso!")
        except sqlite3.IntegrityError:
            st.error("CGM já existe.")
        finally:
            conn.close()

# Página Ocorrências
def pagina_ocorrencias():
    st.title("Registro de Ocorrências 📋")

    cgm_busca = st.text_input("Digite o CGM do aluno")
    nome = telefone = ''
    if cgm_busca:
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute("SELECT nome, telefone FROM alunos WHERE cgm=?", (cgm_busca,))
        resultado = cursor.fetchone()
        conn.close()

        if resultado:
            nome, telefone = resultado
            st.write(f"Nome: {nome}")
            st.write(f"Telefone: {telefone}")
        else:
            st.warning("Aluno não encontrado.")

    ocorrencia_texto = st.text_area("Descrição da Ocorrência")
    if st.button("Salvar Ocorrência"):
        if cgm_busca and ocorrencia_texto:
            data_atual = datetime.now().strftime("%Y-%m-%d %H:%M")
            conn = conectar()
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO ocorrencias (cgm, nome, telefone, data, descricao)
                VALUES (?, ?, ?, ?, ?)
            """, (cgm_busca, nome, telefone, data_atual, ocorrencia_texto))
            conn.commit()
            conn.close()
            st.success("Ocorrência registrada com sucesso!")
            st.experimental_rerun()
        else:
            st.warning("Preencha todos os campos.")

    st.subheader("Ocorrências Cadastradas")
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute("SELECT id, cgm, nome, data, descricao FROM ocorrencias")
    todas = cursor.fetchall()
    conn.close()

    for id, cgm, nome, data, desc in todas:
        with st.expander(f"{id} - {nome} ({cgm}) em {data}"):
            st.write(desc)
            col1, col2 = st.columns(2)
            with col1:
                if st.button(f"Excluir {id}"):
                    conn = conectar()
                    cursor = conn.cursor()
                    cursor.execute("DELETE FROM ocorrencias WHERE id=?", (id,))
                    conn.commit()
                    conn.close()
                    st.experimental_rerun()

# Página Cadastro de Usuário
def pagina_cadastro_usuario():
    st.title("Cadastro de Usuário 🧑‍💼")

    with st.form("form_cadastro_usuario"):
        nome = st.text_input("Nome completo")
        usuario = st.text_input("Nome de usuário")
        senha = st.text_input("Senha", type="password")
        setor = st.text_input("Setor")

        submit = st.form_submit_button("Cadastrar")

        if submit:
            if nome and usuario and senha and setor:
                conn = conectar()
                cursor = conn.cursor()
                try:
                    cursor.execute("""
                        INSERT INTO usuarios (nome, usuario, senha, setor)
                        VALUES (?, ?, ?, ?)
                    """, (nome, usuario, senha, setor))
                    conn.commit()
                    st.success(f"Usuário '{usuario}' cadastrado com sucesso!")
                except sqlite3.IntegrityError:
                    st.error("Usuário já existe.")
                finally:
                    conn.close()
            else:
                st.warning("Preencha todos os campos.")

# Página Lista de Alunos
def pagina_lista_alunos():
    st.title("Lista de Alunos Cadastrados")

    conn = conectar()
    cursor = conn.cursor()
    cursor.execute("SELECT cgm, nome, telefone FROM alunos")
    alunos = cursor.fetchall()
    conn.close()

    if alunos:
        st.dataframe(alunos, use_container_width=True)
    else:
        st.warning("Nenhum aluno cadastrado ainda.")

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import os

def pagina_relatorios():
    st.title("Relatório de Ocorrências 📄")

    data_inicio = st.date_input("Data Inicial")
    data_fim = st.date_input("Data Final")

    if st.button("Gerar Relatório"):
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT cgm, nome, data, descricao FROM ocorrencias
            WHERE date(data) BETWEEN ? AND ?
        """, (str(data_inicio), str(data_fim)))
        resultados = cursor.fetchall()
        conn.close()

        if resultados:
            st.success(f"{len(resultados)} ocorrências encontradas.")

            # Exportar para Word
            doc = Document()
            doc.add_heading("Relatório de Ocorrências", 0)
            for cgm, nome, data, desc in resultados:
                doc.add_paragraph(f"CGM: {cgm}\nNome: {nome}\nData: {data}\nDescrição: {desc}\n----------------------")
            doc_path = "relatorio_ocorrencias.docx"
            doc.save(doc_path)

            # Exportar para PDF
            pdf_path = "relatorio_ocorrencias.pdf"
            c = canvas.Canvas(pdf_path, pagesize=A4)
            width, height = A4
            y = height - 50
            c.setFont("Helvetica", 12)
            c.drawString(50, y, "Relatório de Ocorrências")
            y -= 30

            for cgm, nome, data, desc in resultados:
                texto = f"CGM: {cgm} | Nome: {nome} | Data: {data}"
                c.drawString(50, y, texto)
                y -= 20
                for linha in desc.split("\n"):
                    c.drawString(60, y, linha)
                    y -= 15
                y -= 10
                if y < 50:
                    c.showPage()
                    y = height - 50

            c.save()

            # Links de download
            with open(doc_path, "rb") as file:
                st.download_button("📥 Baixar Word (.docx)", file, file_name="relatorio_ocorrencias.docx")

            with open(pdf_path, "rb") as file:
                st.download_button("📥 Baixar PDF (.pdf)", file, file_name="relatorio_ocorrencias.pdf")

        else:
            st.warning("Nenhuma ocorrência encontrada no período selecionado.")


# Menu principal
def menu_principal():
    st.sidebar.image("BRASÃO.png", width=200)  # Sua logo
    menu = ["Cadastro de Alunos", "Ocorrências", "Cadastro de Usuário", "Lista de Alunos", "Relatórios de Ocorrências"]
    escolha = st.sidebar.selectbox("Menu", menu)

    if escolha == "Cadastro de Alunos":
        pagina_cadastro_aluno()
    elif escolha == "Ocorrências":
        pagina_ocorrencias()
    elif escolha == "Cadastro de Usuário":
        pagina_cadastro_usuario()
    elif escolha == "Lista de Alunos":
        pagina_lista_alunos()
    elif escolha == "Relatórios de Ocorrências":
        pagina_relatorios()

# Execução
if "logado" not in st.session_state:
    st.session_state["logado"] = False

criar_tabelas()

if not st.session_state["logado"]:
    login()
else:
    menu_principal()

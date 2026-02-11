import streamlit as st
from pymongo import MongoClient
from datetime import datetime
from bson.objectid import ObjectId
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import pandas as pd
import urllib.parse
import uuid
import pytz

import base64

st.set_page_config(page_title="Sistema Escolar - CCMLC by Leandro Malheiros V2.0.3 ", layout="centered")

# --- Estiliza√ß√£o Visual ---
st.markdown("""
    <style>
        .stApp {
            background-color: #f2f6fc;
            color: #333333;
            font-family: 'Segoe UI', sans-serif;
        }
        h1, h2, h3 {
            color: #003366;
        }
        .block-container {
            max-width: 1000px;
            margin: auto;
            padding: 2rem;
            background-color: white;
            border-radius: 15px;
            box-shadow: 2px 2px 15px rgba(0,0,0,0.1);
        }
        .css-1d391kg {
            padding: 2rem 1rem 2rem 1rem;
        }
        .stButton>button {
            background-color: #003366;
            color: white;
            border-radius: 8px;
            padding: 0.5rem 1rem;
            font-weight: bold;
        }
        .stButton>button:hover {
            background-color: #0055a5;
            transition: 0.3s;
        }
        .stSelectbox, .stTextInput, .stTextArea {
            background-color: #e9f0fa;
            border-radius: 8px;
        }
        .stMarkdown {
            font-size: 1.1rem;
        }
    </style>
""", unsafe_allow_html=True)

def agora_local():
    tz = pytz.timezone("America/Sao_Paulo")
    return datetime.now(tz)
    
# --- Conex√£o com MongoDB ---
@st.cache_resource
def conectar():
    uri = "mongodb+srv://bibliotecaluizcarlos:KAUOQ9ViyKrXDDAl@cluster0.npyoxsi.mongodb.net/?retryWrites=true&w=majority"
    cliente = MongoClient(uri)
    return cliente["escola"]

db = conectar()

print("--- Cole√ß√µes no banco 'escola' ---")
print(db.list_collection_names())

# --- Fun√ß√µes auxiliares ---
from datetime import datetime
import pandas as pd

def data_segura(valor):
    try:
        if not valor:
            return datetime.now().date()

        data = pd.to_datetime(valor, errors="coerce")

        if pd.isna(data):
            return datetime.now().date()

        return data.date()
    except Exception:
        return datetime.now().date()
def formatar_mensagem_whatsapp(ocorrencias, nome):
    msg = f"""üìã RELAT√ìRIO DE OCORR√äNCIAS
üë§ Aluno: {nome}
üìÖ Data do Relat√≥rio: {datetime.now().strftime('%d/%m/%y √†s %H:%M')}
==============================\n"""

    for i, ocorr in enumerate(ocorrencias, start=1):
        data_txt = ocorr.get("data", "")
        data_formatada = data_txt
        if data_txt:
            for fmt in ("%d-%m-%Y %H:%M:%S", "%d-%m-%Y %H:%M"):
                try:
                    data_obj = datetime.strptime(data_txt, fmt)
                    data_formatada = data_obj.strftime("%Y/%m/%d √†s %H:%M")
                    break
                except ValueError:
                    continue
        msg += f"""
üî∏ Ocorr√™ncia {i}
üìÖ Data: {data_formatada}
üìù Descri√ß√£o: {ocorr['descricao']}
-------------------------"""

    msg += """

üë®‚Äçüè´ Escola [CCM Prof¬∫ Luiz Carlos de Paula e Souza]
üìû Contato: [41 3348-4165]

Este relat√≥rio foi gerado automaticamente pelo Sistema de Ocorr√™ncias."""
    return msg

def exportar_ocorrencias_para_word(ocorrencias, nome_arquivo):
    import os
    import base64
    from docx import Document

    caminho = os.path.join(os.getcwd(), nome_arquivo)

    pasta_ata = os.path.join(os.getcwd(), "atas_tmp")
    os.makedirs(pasta_ata, exist_ok=True)

    doc = Document()
    doc.add_heading("RELAT√ìRIO DE OCORR√äNCIAS", level=1)

    for i, o in enumerate(ocorrencias, start=1):
        doc.add_paragraph(f"Aluno: {o.get('nome', '')}")
        doc.add_paragraph(f"CGM: {o.get('cgm', '')}")
        doc.add_paragraph(f"Data: {o.get('data', '')}")
        doc.add_paragraph(f"Descri√ß√£o: {o.get('descricao', '')}")

        ata = o.get("ata")

        # =========================
        # ATA COMO ARQUIVO
        # =========================
        if isinstance(ata, dict):
            try:
                nome_ata = ata.get("nome", f"ATA_{i}")
                conteudo = ata.get("conteudo")

                if conteudo:
                    bytes_ata = base64.b64decode(conteudo)
                    caminho_ata = os.path.join(pasta_ata, nome_ata)

                    with open(caminho_ata, "wb") as f:
                        f.write(bytes_ata)

                    doc.add_paragraph(f"ATA anexada: {nome_ata}")

            except Exception as e:
                doc.add_paragraph("ATA inv√°lida ou corrompida.")

        # =========================
        # ATA ANTIGA (TEXTO)
        # =========================
        elif isinstance(ata, str) and ata.strip():
            doc.add_paragraph(f"ATA (texto): {ata}")

        doc.add_paragraph("-" * 40)

    doc.save(caminho)
    return caminho

def exportar_ocorrencias_para_pdf(ocorrencias, nome_arquivo):
    import os
    import base64
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    caminho = os.path.join(os.getcwd(), nome_arquivo)
    c = canvas.Canvas(caminho, pagesize=A4)
    largura, altura = A4

    y = altura - 40
    c.setFont("Helvetica", 10)
    c.drawString(40, y, "RELAT√ìRIO DE OCORR√äNCIAS")
    y -= 30

    for o in ocorrencias:
        linhas = [
            f"Aluno: {o.get('nome', '')}",
            f"CGM: {o.get('cgm', '')}",
            f"Data: {o.get('data', '')}",
            f"Descri√ß√£o: {o.get('descricao', '')}",
        ]

        ata = o.get("ata", "")
        if ata:
            try:
                base64.b64decode(ata)
                linhas.append("ATA: arquivo anexado (PDF/JPG)")
            except Exception:
                linhas.append(f"ATA (texto): {ata}")

        linhas.append("-" * 60)

        for linha in linhas:
            if y < 50:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = altura - 40
            c.drawString(40, y, linha)
            y -= 15

    c.save()
    return caminho


# --- Login ---
def pagina_login():
    st.markdown("## üë§ Login de Usu√°rio - V2.0.3 LSM")
    usuario = st.text_input("Usu√°rio").strip()
    senha = st.text_input("Senha", type="password").strip()

    if st.button("Entrar"):
        senha_hash = hashlib.sha256(senha.encode()).hexdigest()
        user = db.usuarios.find_one({
            "usuario": usuario,
            "senha": senha_hash
        })

        if user:
            st.session_state["logado"] = True
            st.session_state["usuario"] = usuario
            st.session_state["nivel"] = user.get("nivel", "user")
            st.success("‚úÖ Login realizado com sucesso!")
            st.rerun()
        else:
            st.error("Usu√°rio ou senha inv√°lidos.")

# --- Cadastro de Alunos ---
def pagina_cadastro():
    st.markdown("## ‚úèÔ∏è Cadastro de Alunos")

    # --- Lista de alunos cadastrados ---
    alunos = list(db.alunos.find().sort("nome", 1))

    nomes_exibicao = [""] + [
        f"{a['nome']} (CGM: {a['cgm']})"
        for a in alunos
    ]

    selecionado = st.selectbox("üîé Buscar aluno para Alterar ou Excluir:", nomes_exibicao)

    aluno_carregado = None
    if selecionado and selecionado != "":
        # Extrai CGM do texto selecionado
        cgm_busca = selecionado.split("CGM:")[1].replace(")", "").strip()
        aluno_carregado = db.alunos.find_one({"cgm": cgm_busca})

        st.success(f"Aluno carregado: {aluno_carregado['nome']} (CGM {aluno_carregado['cgm']})")

    # --- Formul√°rio de Cadastro ou Altera√ß√£o ---
    with st.form("form_cadastro"):

        cgm = st.text_input("CGM", value=aluno_carregado["cgm"] if aluno_carregado else "")
        nome = st.text_input("Nome", value=aluno_carregado["nome"] if aluno_carregado else "")
        data = st.date_input(
        "Data de Nascimento",
            value=data_segura(aluno_carregado.get("data") if aluno_carregado else None)
    )
        telefone = st.text_input("Telefone", value=aluno_carregado["telefone"] if aluno_carregado else "")
        responsavel = st.text_input("Respons√°vel", value=aluno_carregado["responsavel"] if aluno_carregado else "")
        turma = st.text_input("Turma", value=aluno_carregado["turma"] if aluno_carregado else "")

        col1, col2, col3 = st.columns([1,1,1])
        salvar = col1.form_submit_button("üíæ Salvar / Alterar")
        excluir = col2.form_submit_button("üóëÔ∏è Excluir")
        limpar = col3.form_submit_button("üßπ Limpar")

    # --- A√ß√µes ap√≥s clique ---
    if salvar:
        if cgm and nome:
            db.alunos.update_one({"cgm": cgm}, {
                "$set": {
                    "cgm": cgm,
                    "nome": nome,
                    "data": str(data),
                    "telefone": telefone,
                    "responsavel": responsavel,
                    "turma": turma
                }
            }, upsert=True)
            st.success("‚úÖ Aluno salvo ou atualizado com sucesso!")
            st.experimental_rerun()
        else:
            st.error("Preencha todos os campos obrigat√≥rios.")

    if excluir and aluno_carregado:
        confirmacao = st.warning(f"Tem certeza que deseja excluir o aluno {aluno_carregado['nome']} (CGM {aluno_carregado['cgm']})?")
        if st.button("‚úÖ Confirmar Exclus√£o"):
            db.alunos.delete_one({"cgm": aluno_carregado["cgm"]})
            st.success("‚úÖ Aluno exclu√≠do com sucesso!")
            st.experimental_rerun()

    if limpar:
        st.experimental_rerun()

    # --- Importa√ß√£o de alunos via arquivo ---
    st.subheader("üì• Importar Alunos via TXT ou CSV")
    arquivo = st.file_uploader("Escolha o arquivo .txt ou .csv", type=["txt", "csv"])
    delimitador = st.selectbox("Escolha o delimitador", [";", ",", "\\t"])
    delimitador_real = {";": ";", ",": ",", "\\t": "\t"}[delimitador]

    if arquivo is not None:
        try:
            df_import = pd.read_csv(arquivo, delimiter=delimitador_real)
            df_import.columns = [col.strip().lower() for col in df_import.columns]
            st.dataframe(df_import)

            if st.button("Importar para o Sistema"):
                erros = []
                total_importados = 0
                for _, row in df_import.iterrows():
                    try:
                        cgm = str(row.get('cgm', '')).strip()
                        nome = str(row.get('nome', '')).strip()
                        data = str(row.get('data', '')).strip()
                        telefone = str(row.get('telefone', '')).strip()
                        responsavel = str(row.get('responsavel', '')).strip()
                        turma = str(row.get('turma', '')).strip()

                        if not cgm or not nome:
                            erros.append(f"CGM ou Nome ausente na linha: {row.to_dict()}")
                            continue

                        aluno = {
                            "cgm": cgm,
                            "nome": nome,
                            "data": data,
                            "telefone": telefone,
                            "responsavel": responsavel,
                            "turma": turma
                        }

                        db.alunos.update_one({"cgm": cgm}, {"$set": aluno}, upsert=True)
                        total_importados += 1

                    except Exception as e:
                        erros.append(f"Erro na linha {row.to_dict()} ‚Üí {e}")

                st.success(f"‚úÖ Importa√ß√£o finalizada. Total importado/atualizado: {total_importados}")
                if erros:
                    st.warning("‚ö†Ô∏è Erros encontrados:")
                    for erro in erros:
                        st.error(erro)

        except Exception as e:
            st.error(f"Erro ao ler o arquivo: {e}")

def pagina_ocorrencias():
    st.markdown("## üö® Registro de Ocorr√™ncia")

    alunos = list(db.alunos.find())
    alunos_ordenados = sorted(alunos, key=lambda x: x['nome'])

    busca_cgm = st.text_input("üîç Buscar aluno por CGM")

    if busca_cgm:
        aluno_cgm = next((a for a in alunos_ordenados if a["cgm"] == busca_cgm), None)
        if aluno_cgm:
            nomes = [f"{aluno_cgm['nome']} (CGM: {aluno_cgm['cgm']})"]
        else:
            st.warning("Nenhum aluno encontrado com esse CGM.")
            return
    else:
        nomes = [""] + [f"{a['nome']} (CGM: {a['cgm']})" for a in alunos_ordenados]

    if nomes:
        selecionado = st.selectbox("Selecione o aluno:", nomes)

        if selecionado != "":
            cgm = selecionado.split("CGM: ")[1].replace(")", "")
            nome = selecionado.split(" (CGM:")[0]

            ocorrencias = list(db.ocorrencias.find({"cgm": cgm}))
            opcoes_ocorrencias = ["Nova Ocorr√™ncia"] + [
                f"{o['data']} - {o['descricao'][:30]}..." for o in ocorrencias
            ]

            ocorrencia_selecionada = st.selectbox("üìå Ocorr√™ncia:", opcoes_ocorrencias)

            descricao = ""
            ata = ""

            # ================= NOVA OCORR√äNCIA =================
            if ocorrencia_selecionada == "Nova Ocorr√™ncia":
                descricao = st.text_area("‚úèÔ∏è Descri√ß√£o da Ocorr√™ncia", key="descricao_nova")
                ata = st.text_input("üìÑ ATA (opcional)", key="ata_nova")

                arquivo_ata = st.file_uploader(
                    "üì§ Importar ATA (PDF ou JPG)",
                    type=["pdf", "jpg", "jpeg"],
                    key="upload_ata_nova"
                )

                if arquivo_ata:
                    ata = base64.b64encode(arquivo_ata.read()).decode("utf-8")

                if st.button("‚úÖ Registrar Nova Ocorr√™ncia", key="btn_nova") and descricao:
                    agora = agora_local().strftime("%Y-%m-%d %H:%M:%S")
                    telefone = next((a['telefone'] for a in alunos if a['cgm'] == cgm), "")

                    db.ocorrencias.insert_one({
                        "cgm": cgm,
                        "nome": nome,
                        "telefone": telefone,
                        "data": agora,
                        "descricao": descricao,
                        "ata": ata
                    })

                    st.success("‚úÖ Ocorr√™ncia registrada com sucesso!")

            # ================= OCORR√äNCIA EXISTENTE =================
            else:
                index = opcoes_ocorrencias.index(ocorrencia_selecionada) - 1
                ocorrencia = ocorrencias[index]

                descricao = st.text_area(
                    "‚úèÔ∏è Descri√ß√£o da Ocorr√™ncia",
                    value=ocorrencia.get("descricao", ""),
                    key=f"desc_{ocorrencia['_id']}"
                )

                ata = st.text_input(
                    "üìÑ ATA (opcional)",
                    value=ocorrencia.get("ata", ""),
                    key=f"ata_{ocorrencia['_id']}"
                )

                arquivo_ata = st.file_uploader(
                    "üì§ Importar nova ATA (PDF ou JPG)",
                    type=["pdf", "jpg", "jpeg"],
                    key=f"upload_ata_{ocorrencia['_id']}"
                )

                if arquivo_ata:
                    ata = base64.b64encode(arquivo_ata.read()).decode("utf-8")

                col1, col2 = st.columns(2)

                with col1:
                    if st.button("üíæ Alterar Ocorr√™ncia", key=f"alt_{ocorrencia['_id']}"):
                        db.ocorrencias.update_one(
                            {"_id": ocorrencia["_id"]},
                            {"$set": {
                                "descricao": descricao,
                                "ata": ata
                            }}
                        )
                        st.success("‚úÖ Ocorr√™ncia atualizada com sucesso!")

                with col2:
                    confirmar_exclusao = st.checkbox(
                        "Confirmar exclus√£o",
                        key=f"conf_{ocorrencia['_id']}"
                    )
                    if confirmar_exclusao:
                        if st.button("üóëÔ∏è Excluir Ocorr√™ncia", key=f"del_{ocorrencia['_id']}"):
                            db.ocorrencias.delete_one({"_id": ocorrencia["_id"]})
                            st.success("üóëÔ∏è Ocorr√™ncia exclu√≠da com sucesso!")
                            st.experimental_rerun()

def pagina_exportar():
    import urllib
    import uuid

    st.markdown("## üì• Exportar Relat√≥rios")

    resultados = list(db.ocorrencias.find({}))
    if not resultados:
        st.warning("Nenhuma ocorr√™ncia encontrada.")
        return

    # ===================== BUSCA POR CGM =====================
    st.subheader("üîç Buscar por CGM")
    cgm_input = st.text_input("Digite o CGM do aluno")
    col1, col2 = st.columns(2)

    if col1.button("üìÑ Gerar Word por CGM", key="word_cgm") and cgm_input:
        dados = list(db.ocorrencias.find({"cgm": cgm_input}))
        if dados:
            caminho = exportar_ocorrencias_para_word(dados, f"ocorrencias_{cgm_input}.docx")
            with open(caminho, "rb") as f:
                st.download_button(
                    "üì• Baixar Word",
                    f.read(),
                    file_name=f"ocorrencias_{cgm_input}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    if col2.button("üßæ Gerar PDF por CGM", key="pdf_cgm") and cgm_input:
        dados = list(db.ocorrencias.find({"cgm": cgm_input}))
        if dados:
            caminho = exportar_ocorrencias_para_pdf(dados, f"ocorrencias_{cgm_input}.pdf")
            with open(caminho, "rb") as f:
                st.download_button(
                    "üì• Baixar PDF",
                    f.read(),
                    file_name=f"ocorrencias_{cgm_input}.pdf",
                    mime="application/pdf"
                )

    # ===================== PER√çODO =====================
    st.subheader("üìÖ Exportar por Per√≠odo")
    uid = str(uuid.uuid4())
    data_inicio = st.date_input("Data inicial", key=f"ini_{uid}")
    data_fim = st.date_input("Data final", key=f"fim_{uid}")

    if st.button("üîé Gerar relat√≥rio por per√≠odo", key=f"periodo_{uid}"):
        inicio = data_inicio.strftime("%Y-%m-%d")
        fim = data_fim.strftime("%Y-%m-%d") + " 23:59:59"

        dados = list(db.ocorrencias.find({"data": {"$gte": inicio, "$lte": fim}}))
        if dados:
            caminho = exportar_ocorrencias_para_word(dados, "relatorio_periodo.docx")
            with open(caminho, "rb") as f:
                st.download_button(
                    "üì• Baixar DOCX",
                    f.read(),
                    file_name="relatorio_periodo.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            caminho_pdf = exportar_ocorrencias_para_pdf(dados, "relatorio_periodo.pdf")
            with open(caminho_pdf, "rb") as f:
                st.download_button(
                    "üì• Baixar PDF",
                    f.read(),
                    file_name="relatorio_periodo.pdf",
                    mime="application/pdf"
                )

    # ===================== AGRUPADO POR ALUNO =====================
    st.subheader("üìÑ Relat√≥rios Individuais por Aluno")

    ocorrencias_por_aluno = {}
    for ocorr in resultados:
        nome = ocorr.get("nome", "")
        ocorrencias_por_aluno.setdefault(nome, []).append(ocorr)

    for nome, lista in sorted(ocorrencias_por_aluno.items()):
        with st.expander(f"üìÑ Relat√≥rio de {nome}"):
            telefone = lista[0].get("telefone", "")

            for ocorr in lista:
                st.write(f"üìÖ {ocorr.get('data', '')} - üìù {ocorr.get('descricao', '')}")

            mensagem = formatar_mensagem_whatsapp(lista, nome)
            st.text_area("üìã WhatsApp", mensagem, height=200, key=f"msg_{nome}_{lista[0]['_id']}")

            if telefone:
                numero = telefone.replace("(", "").replace(")", "").replace("-", "").replace(" ", "")
                link = f"https://api.whatsapp.com/send?phone=55{numero}&text={urllib.parse.quote(mensagem)}"
                st.markdown(f"[üì± Enviar para {telefone}]({link})")

            col1, col2 = st.columns(2)

            if col1.button("üìÑ Gerar DOCX", key=f"doc_{nome}_{lista[0]['_id']}"):
                caminho = exportar_ocorrencias_para_word(lista, f"relatorio_{nome.replace(' ','_')}.docx")
                with open(caminho, "rb") as f:
                    st.download_button(
                        "üì• Baixar DOCX",
                        f.read(),
                        file_name=f"relatorio_{nome.replace(' ','_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            if col2.button("üßæ Gerar PDF", key=f"pdf_{nome}_{lista[0]['_id']}"):
                caminho = exportar_ocorrencias_para_pdf(lista, f"relatorio_{nome.replace(' ','_')}.pdf")
                with open(caminho, "rb") as f:
                    st.download_button(
                        "üì• Baixar PDF",
                        f.read(),
                        file_name=f"relatorio_{nome.replace(' ','_')}.pdf",
                        mime="application/pdf"
                    )

# --- Lista de Alunos ---
def pagina_lista():
    st.markdown("## üìÑ Lista de Alunos")
    dados = list(db.alunos.find({}, {"_id": 0}))
    if dados:
        df = pd.DataFrame(dados)
        st.dataframe(df.sort_values("nome"))
    else:
        st.info("Nenhum aluno cadastrado.")

# --- Cadastro de Usu√°rios ---
import streamlit as st
from pymongo import MongoClient
import hashlib

def conectar():
    uri = "mongodb+srv://bibliotecaluizcarlos:KAUOQ9ViyKrXDDAl@cluster0.npyoxsi.mongodb.net/?retryWrites=true&w=majority"
    cliente = MongoClient(uri)
    return cliente["escola"]

db = conectar()

# --- Cadastro de Usu√°rios ---
def pagina_usuarios():
    st.markdown("## üë• Cadastro de Usu√°rios")
    
    # Exemplo de seguran√ßa: s√≥ admin pode cadastrar
    if st.session_state.get("nivel") != "admin":
        st.warning("Apenas administradores podem cadastrar novos usu√°rios.")
        return

    # Formul√°rio de cadastro
    with st.form("form_usuarios"):
        usuario = st.text_input("Novo usu√°rio")
        senha = st.text_input("Senha", type="password")
        nivel = st.selectbox("N√≠vel de acesso", ["user", "admin"])
        cadastrar = st.form_submit_button("Cadastrar")

    if cadastrar:
        usuario = usuario.strip()
        senha = senha.strip()
        if usuario and senha:
            senha_hash = hashlib.sha256(senha.encode()).hexdigest()
            try:
                resultado = db.usuarios.insert_one({
                    "usuario": usuario,
                    "senha": senha_hash,
                    "nivel": nivel
                })
                st.success("‚úÖ Usu√°rio cadastrado com sucesso!")
                print("Usu√°rio salvo com id:", resultado.inserted_id)
            except Exception as e:
                print("Erro ao salvar usu√°rio:", e)
                st.error(f"Erro ao salvar usu√°rio: {e}")
        else:
            st.error("Preencha todos os campos.")

    if st.button("üëÄ Ver Usu√°rios Salvos"):
        usuarios = list(db.usuarios.find())
        if usuarios:
            for u in usuarios:
                st.write(u)
        else:
            st.info("Nenhum usu√°rio cadastrado ainda.")

# --- Menu Lateral ---
def menu():
    st.sidebar.image("BRAS√ÉO.png", use_container_width=True)
    st.sidebar.markdown("### üìö Menu de Navega√ß√£o")
    opcoes = ["Cadastro", "Ocorr√™ncias", "Exportar", "Lista"]
    if st.session_state.get("nivel") == "admin":
        opcoes.append("Usu√°rios")
    pagina = st.sidebar.selectbox("Escolha a aba:", opcoes)

    if pagina == "Cadastro":
        pagina_cadastro()
    elif pagina == "Ocorr√™ncias":
        pagina_ocorrencias()
    elif pagina == "Exportar":
        pagina_exportar()
    elif pagina == "Lista":
        pagina_lista()
    elif pagina == "Usu√°rios":
        pagina_usuarios()

def sair():
    st.session_state["logado"] = False
    st.session_state["usuario"] = ""
    st.session_state["nivel"] = ""
    st.rerun()
    
# --- Execu√ß√£o ---
if "logado" not in st.session_state:
    st.session_state["logado"] = False

if not st.session_state["logado"]:
    pagina_login()
else:
    if st.sidebar.button("üö™ Sair do Sistema"):
        sair()
    else:
        menu()

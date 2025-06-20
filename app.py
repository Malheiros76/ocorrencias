import streamlit as st
import sqlite3
from datetime import datetime
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import os

def conectar():
    return sqlite3.connect('ocorrencias.db')

def criar_tabelas():
    conn = conectar()
    cursor = conn.cursor()

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS usuarios (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        nome TEXT NOT NULL,
        usuario TEXT UNIQUE NOT NULL,
        senha TEXT NOT NULL,
        setor TEXT NOT NULL
    )
    """)
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS alunos (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        cgm TEXT UNIQUE NOT NULL,
        nome TEXT NOT NULL,
        data_nascimento TEXT,
        telefone TEXT,
        responsavel TEXT,
        data TEXT,
        turma TEXT
    )
    """)
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS ocorrencias (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        cgm TEXT NOT NULL,
        nome TEXT,
        telefone TEXT,
        data TEXT,
        descricao TEXT
    )
    """)
    conn.commit()
    conn.close()

def login():
    st.title("Login 👤")
    usuario = st.text_input("Usuário")
    senha = st.text_input("Senha", type="password")
    if st.button("Entrar"):
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM usuarios WHERE usuario=? AND senha=?", (usuario, senha))
        resultado = cursor.fetchone()
        conn.close()
        if resultado:
            st.session_state["logado"] = True
            st.rerun()
        else:
            st.error("Usuário ou senha inválidos!")

def pagina_cadastro_aluno():
    st.title("Cadastro de Alunos 👦👧")
    arquivo = st.file_uploader("Importar lista de alunos (.txt)", type="txt")
    if arquivo:
        conteudo = arquivo.read().decode('utf-8')
        linhas = conteudo.strip().split('\n')[1:]
        dados_alunos = []
        for linha in linhas:
            campos = linha.split('\t')
            if len(campos) >= 3:
                cgm, nome, telefone = campos[0], campos[1], campos[2]
                dados_alunos.append((cgm, nome, telefone))

        if dados_alunos:
            st.dataframe(dados_alunos)
            if st.button("Salvar no banco"):
                conn = conectar()
                cursor = conn.cursor()
                for aluno in dados_alunos:
                    try:
                        cursor.execute("INSERT INTO alunos (cgm, nome, telefone) VALUES (?, ?, ?)", aluno)
                    except sqlite3.IntegrityError:
                        pass
                conn.commit()
                conn.close()
                st.success("Alunos salvos com sucesso!")

    st.subheader("Cadastrar aluno manualmente")
    cgm = st.text_input("CGM")
    nome = st.text_input("Nome")
    data_nascimento = st.date_input("Data de Nascimento")
    telefone = st.text_input("Telefone")
    responsavel = st.text_input("Responsável")
    data = st.date_input("Data de Cadastro")
    turma = st.text_input("Turma")
    if st.button("Cadastrar Aluno"):
        conn = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                INSERT INTO alunos (cgm, nome, data_nascimento, telefone, responsavel, data, turma)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (cgm, nome, str(data_nascimento), telefone, responsavel, str(data), turma))
            conn.commit()
            st.success("Aluno cadastrado com sucesso!")
        except sqlite3.IntegrityError:
            st.error("CGM já existe.")
        finally:
            conn.close()

def exportar_ocorrencias_para_word(resultados):
    doc = Document()
    doc.add_picture("CABEÇARIOAPP.png", width=Inches(6))
    doc.add_heading("Relatório de Ocorrências", level=1)
    for cgm, nome, data, desc in resultados:
        doc.add_paragraph(f"CGM: {cgm}\nNome: {nome}\nData: {data}\nDescrição: {desc}\n----------------------")
    doc.add_paragraph("\n\nAssinatura do Servidor: ____________________________")
    doc.add_paragraph("\nAssinatura do Responsável: ____________________________")
    doc.add_paragraph("\nData: _______/_______/_________")
    doc_path = "relatorio_ocorrencias.docx"
    doc.save(doc_path)
    return doc_path

from fpdf import FPDF

def exportar_ocorrencias_para_pdf(resultados):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    
    # Cabeçalho com imagem
    pdf.image("CABEÇARIOAPP.png", x=10, y=8, w=190)
    pdf.ln(35)  # Espaço após a imagem
    
    pdf.cell(0, 10, "Relatório de Ocorrências", ln=True, align='C')
    pdf.ln(5)
    
    pdf.set_font("Arial", size=12)

    for cgm, nome, data, desc in resultados:
        pdf.multi_cell(0, 8, f"CGM: {cgm}\nNome: {nome}\nData: {data}\nDescrição: {desc}")
        pdf.ln(2)
        pdf.cell(0, 0, "-" * 70, ln=True)  # Linha de separação
        pdf.ln(5)

    # Espaço para assinaturas
    pdf.ln(10)
    pdf.cell(0, 10, "Assinatura do Servidor: ____________________________", ln=True)
    pdf.cell(0, 10, "Assinatura do Responsável: _________________________", ln=True)
    pdf.cell(0, 10, "Data: ______/______/________", ln=True)
    
    caminho = "relatorio_ocorrencias.pdf"
    pdf.output(caminho)
    return caminho

def pagina_ocorrencias():
    st.title("Registro de Ocorrências 📋")
    cgm_busca = st.text_input("Digite o CGM do aluno")
    nome, telefone = "", ""

    if cgm_busca:
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute("SELECT nome, telefone FROM alunos WHERE cgm=?", (cgm_busca,))
        resultado = cursor.fetchone()
        conn.close()
        if resultado:
            nome, telefone = resultado
            st.write(f"Nome: {nome}")
            st.write(f"Telefone: {telefone}")
        else:
            st.warning("Aluno não encontrado.")

    ocorrencia_texto = st.text_area("Descrição da Ocorrência")
    if st.button("Salvar Ocorrência"):
        if cgm_busca and ocorrencia_texto:
            data_atual = datetime.now().strftime("%Y-%m-%d %H:%M")
            conn = conectar()
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO ocorrencias (cgm, nome, telefone, data, descricao)
                VALUES (?, ?, ?, ?, ?)
            """, (cgm_busca, nome, telefone, data_atual, ocorrencia_texto))
            conn.commit()
            conn.close()
            st.success("Ocorrência registrada com sucesso!")
        else:
            st.warning("Preencha todos os campos.")

    st.subheader("Consultar / Exportar Ocorrências")
    filtro_cgm = st.text_input("Filtrar por CGM")
    data_inicio = st.date_input("Data Início")
    data_fim = st.date_input("Data Fim")

    conn = conectar()
    cursor = conn.cursor()
    query = "SELECT cgm, nome, data, descricao FROM ocorrencias WHERE 1=1"
    params = []

    if filtro_cgm:
        query += " AND cgm=?"
        params.append(filtro_cgm)
    if data_inicio and data_fim:
        query += " AND DATE(data) BETWEEN ? AND ?"
        params.append(str(data_inicio))
        params.append(str(data_fim))

    cursor.execute(query, params)
    resultados = cursor.fetchall()
    conn.close()

    if resultados:
        for cgm, nome, data, desc in resultados:
            st.write(f"CGM: {cgm} | Nome: {nome} | Data: {data}\nDescrição: {desc}")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Exportar para Word"):
                caminho = exportar_ocorrencias_para_word(resultados)
                with open(caminho, "rb") as file:
                    st.download_button("Download Word", data=file, file_name="relatorio_ocorrencias.docx")
        with col2:
            if st.button("Exportar para PDF"):
                caminho = exportar_ocorrencias_para_pdf(resultados)
                with open(caminho, "rb") as file:
                    st.download_button("Download PDF", data=file, file_name="relatorio_ocorrencias.pdf")
    else:
        st.warning("Nenhuma ocorrência encontrada para os filtros selecionados.")

def pagina_cadastro_usuario():
    st.title("Cadastro de Usuário 🧑‍💼")
    with st.form("form_cadastro_usuario"):
        nome = st.text_input("Nome completo")
        usuario = st.text_input("Nome de usuário")
        senha = st.text_input("Senha", type="password")
        setor = st.text_input("Setor")
        submit = st.form_submit_button("Cadastrar")
        if submit:
            if nome and usuario and senha and setor:
                conn = conectar()
                cursor = conn.cursor()
                try:
                    cursor.execute("""
                        INSERT INTO usuarios (nome, usuario, senha, setor)
                        VALUES (?, ?, ?, ?)
                    """, (nome, usuario, senha, setor))
                    conn.commit()
                    st.success(f"Usuário '{usuario}' cadastrado com sucesso!")
                except sqlite3.IntegrityError:
                    st.error("Usuário já existe.")
                finally:
                    conn.close()
            else:
                st.warning("Preencha todos os campos.")

# Lista Alunos
def pagina_lista_alunos():
    st.header("Lista de Alunos")
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute("SELECT cgm, nome, telefone FROM alunos ORDER BY nome ASC")
    alunos = cursor.fetchall()
    conn.close()
    if alunos:
        st.dataframe(alunos)
    else:
        st.warning("Nenhum aluno cadastrado.")
        
def menu_principal():
    st.sidebar.image("BRASÃO.png", width=200)
    menu = ["Cadastro de Alunos", "Ocorrências", "Cadastro de Usuário", "Lista de Alunos"]
    escolha = st.sidebar.selectbox("Menu", menu)

    if escolha == "Cadastro de Alunos":
        pagina_cadastro_aluno()
    elif escolha == "Ocorrências":
        pagina_ocorrencias()
    elif escolha == "Cadastro de Usuário":
        pagina_cadastro_usuario()
    elif escolha == "Lista de Alunos":
        pagina_lista_alunos()

# Inicialização
if "logado" not in st.session_state:
    st.session_state["logado"] = False

criar_tabelas()

if not st.session_state["logado"]:
    login()
else:
    menu_principal()
